import os
import json
import time
import re
import mimetypes
import chardet
from groq import Groq
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from django.conf import settings
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from docx import Document

# Initialize Groq client
def get_groq_client():
    api_key = os.environ.get('GROQ_API_KEY')
    if not api_key:
        raise ValueError("GROQ_API_KEY environment variable is not set. Please add it to your environment variables.")
    return Groq(api_key=api_key)

# Simple ChatBot utility
def get_chat_response(query, model_name="llama3-70b-8192"):
    try:
        client = get_groq_client()
        
        # Add a system prompt to improve response quality
        system_prompt = """You are Carmen, the friendly and knowledgeable AI assistant for QuadraNex-AI.
        You provide clear, accurate, and concise information with a helpful and conversational tone.
        Format your responses with proper HTML for better readability, including headings, lists, and highlighting of important information when appropriate.
        If you're unsure about something, acknowledge it rather than making up information.
        You can discuss a wide range of topics including technology, science, arts, history, and more.
        When providing explanations, use analogies and examples to make complex concepts easier to understand."""
        
        # Check if the query is asking for specific formatting
        formatting_keywords = ["format", "html", "style", "code", "markdown", "highlight"]
        should_format = any(keyword in query.lower() for keyword in formatting_keywords)
        
        # Add formatting instructions if needed
        if should_format:
            system_prompt += """
            When formatting code or technical content:
            - Use <pre><code> tags for code blocks
            - Use appropriate syntax highlighting when possible
            - Format lists with proper HTML tags
            - Use <em> and <strong> for emphasis
            """
        
        response = client.chat.completions.create(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": query}
            ],
            model=model_name,
            temperature=0.7,
            max_tokens=2048
        )
        
        chat_response = response.choices[0].message.content
        
        # Add basic formatting if not already present
        if "<" not in chat_response and ">" not in chat_response:
            # First handle code blocks if present
            code_blocks = chat_response.split("```")
            for i in range(len(code_blocks)):
                if i % 2 == 1:  # This is a code block
                    code_blocks[i] = f"<pre><code>{code_blocks[i]}</code></pre>"
            chat_response = "".join(code_blocks)
            
            # Convert markdown-style formatting to HTML
            # Handle nested formatting by processing strong/bold first
            chat_response = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', chat_response)
            chat_response = re.sub(r'\*(.+?)\*', r'<em>\1</em>', chat_response)
            
            # Handle lists
            lines = chat_response.split("\n")
            in_list = False
            formatted_lines = []
            
            for line in lines:
                if line.strip().startswith("- "):
                    if not in_list:
                        formatted_lines.append("<ul>")
                        in_list = True
                    formatted_lines.append(f"<li>{line.strip()[2:]}</li>")
                elif line.strip().startswith("1. ") or line.strip().startswith("* "):
                    if not in_list:
                        formatted_lines.append("<ol>")
                        in_list = True
                    formatted_lines.append(f"<li>{line.strip()[2:]}</li>")
                else:
                    if in_list:
                        formatted_lines.append("</ul>" if "<ul>" in formatted_lines[-2] else "</ol>")
                        in_list = False
                    if line.strip():
                        formatted_lines.append(line)
            
            if in_list:
                formatted_lines.append("</ul>" if "<ul>" in formatted_lines[-2] else "</ol>")
            
            # Format paragraphs, but skip already formatted content
            paragraphs = "\n".join(formatted_lines).split("\n\n")
            formatted_paragraphs = []
            for p in paragraphs:
                if p.strip():
                    if not (p.strip().startswith('<') and p.strip().endswith('>')):  # Skip already formatted content
                        formatted_paragraphs.append(f"<p>{p.strip()}</p>")
                    else:
                        formatted_paragraphs.append(p.strip())
            chat_response = "\n".join(formatted_paragraphs)
        
        return chat_response
    except ValueError as e:
        # Handle missing API key
        return f"<div class='error-message'>Configuration Error: {str(e)}</div>"
    except Exception as e:
        # Handle other errors
        return f"<div class='error-message'>Sorry, I encountered an error: {str(e)}</div>"

# RAG System utility
def process_document_for_rag(file_path, query):
    try:
        client = get_groq_client()
        
        # Read the document content
        try:
            # Get file extension to determine document type
            file_extension = os.path.splitext(file_path)[1].lower()
            
            # Handle different file types appropriately
            if file_extension == '.pdf':
                try:
                    from PyPDF2 import PdfReader
                    reader = PdfReader(file_path)
                    document_content = '\n'.join([page.extract_text() for page in reader.pages if page.extract_text()])
                    if not document_content.strip():
                        return "<div class='error-message'>Could not extract text from PDF file. The file may be scanned or contain only images.</div>"
                except ImportError:
                    return "<div class='error-message'>PyPDF2 library not installed for PDF processing</div>"
                except Exception as e:
                    return f"<div class='error-message'>Error reading PDF: {str(e)}</div>"
            # Handle DOCX files
            elif file_extension in ['.docx', '.doc']:
                try:
                    doc = Document(file_path)
                    document_content = '\n'.join([para.text for para in doc.paragraphs])
                except ImportError:
                    return "<div class='error-message'>python-docx library not installed for DOCX processing</div>"
                except Exception as e:
                    return f"<div class='error-message'>Error reading DOCX: {str(e)}</div>"
            # Handle text files with proper encoding detection
            else:
                try:
                    # First detect the file encoding
                    with open(file_path, 'rb') as file:
                        raw_data = file.read(10000)  # Read a sample to detect encoding
                        result = chardet.detect(raw_data)
                        encoding = result['encoding'] if result['encoding'] else 'utf-8'
                        confidence = result.get('confidence', 0)
                        
                        # If confidence is low, use a more reliable encoding
                        if confidence < 0.7:
                            encoding = 'utf-8'
                    
                    # Try to read with detected encoding
                    try:
                        with open(file_path, 'r', encoding=encoding) as file:
                            document_content = file.read()
                    except UnicodeDecodeError:
                        # First fallback: utf-8 with error replacement
                        try:
                            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                                document_content = file.read()
                        except UnicodeDecodeError:
                            # Second fallback: latin-1 (should handle any byte sequence)
                            with open(file_path, 'r', encoding='latin-1') as file:
                                document_content = file.read()
                except Exception as e:
                    return f"<div class='error-message'>Error reading file: {str(e)}</div>"
            
            # Get file extension to determine document type
            file_extension = os.path.splitext(file_path)[1].lower()
            document_type = "text"
            if file_extension in ['.md', '.markdown']:
                document_type = "markdown"
            elif file_extension in ['.tex']:
                document_type = "LaTeX"
            elif file_extension in ['.html', '.htm']:
                document_type = "HTML"
            elif file_extension in ['.pdf']:
                document_type = "PDF"
            
            # Create a system prompt for the RAG system
            system_prompt = """You are Sirius, an advanced AI document analysis assistant.
            Your task is to provide comprehensive and accurate answers to questions based on the document content.
            If the answer isn't in the document, acknowledge that rather than making up information.
            Format your responses with proper HTML for better readability, including headings, lists, and highlighting of important information.
            Include relevant quotes from the document to support your answers."""
            
            # Create a prompt for the RAG system
            rag_prompt = f"""
            Document Type: {document_type}
            
            Document Content:
            {document_content[:20000]}  # Increased content size limit
            
            User Question: {query}
            
            Please provide a comprehensive answer to the question based solely on the document content.
            Format your response with the following sections:
            1. Direct Answer: A concise answer to the question
            2. Supporting Evidence: Relevant quotes from the document that support your answer
            3. Additional Context: Any additional information from the document that might be helpful
            4. Related Information: Other relevant information from the document that relates to the question
            
            Use proper HTML formatting for better readability.
            """
            
            response = client.chat.completions.create(
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": rag_prompt}
                ],
                model="llama3-70b-8192",
                temperature=0.3,
                max_tokens=3000
            )
            
            rag_result = response.choices[0].message.content
            
            # Format the result with HTML for better display
            formatted_result = f"""
            <div class="rag-result">
                {rag_result}
            </div>
            """
            
            return formatted_result
        except Exception as e:
            return f"<div class='error-message'>Document processing error: {str(e)}</div>"
    except ValueError as e:
        # Handle missing API key
        return f"<div class='error-message'>Configuration Error: {str(e)}</div>"
    except Exception as e:
        # Handle other errors
        return f"<div class='error-message'>Sorry, I encountered an error: {str(e)}</div>"

# Wikipedia Scraper utility
def scrape_wikipedia(article_url):
    try:
        chrome_options = Options()
        chrome_options.add_argument("--headless")
        chrome_options.add_argument("--no-sandbox")
        chrome_options.add_argument("--disable-dev-shm-usage")
        
        # Setup Chrome WebDriver
        driver = webdriver.Chrome(
            service=Service(ChromeDriverManager().install()),
            options=chrome_options
        )
        
        try:
            # Navigate to Wikipedia article
            driver.get(article_url)
            
            # Extract article title
            article_title = driver.find_element(By.ID, "firstHeading").text
            
            # Extract article content
            article_content = driver.find_element(By.ID, "bodyContent").text
            
            # Extract images
            image_elements = driver.find_elements(By.CSS_SELECTOR, ".image img")
            image_urls = [img.get_attribute("src") for img in image_elements[:5]]  # Limit to 5 images
            
            # Extract headings
            heading_elements = driver.find_elements(By.CSS_SELECTOR, "#bodyContent h2, #bodyContent h3")
            headings = [heading.text.replace("[edit]", "").strip() for heading in heading_elements]
            
            # Format content with headings as sections
            formatted_content = ""
            content_sections = driver.find_elements(By.CSS_SELECTOR, "#bodyContent > *")
            for section in content_sections:
                if section.tag_name in ["h2", "h3", "h4"]:
                    section_title = section.text.replace("[edit]", "").strip()
                    section_id = section_title.lower().replace(" ", "-")
                    formatted_content += f'<h3 id="section-{section_id}">{section_title}</h3>\n'
                elif section.tag_name in ["p", "ul", "ol", "div"]:
                    formatted_content += f'<div class="wiki-section-content">{section.get_attribute("innerHTML")}</div>\n'
            
            # Generate a summary using Groq API
            try:
                client = get_groq_client()
                
                summary_prompt = f"""
                Please provide a concise summary of this Wikipedia article about {article_title}.
                The summary should be about 3-4 paragraphs and highlight the most important information.
                Format the summary with HTML for better readability.
                
                Article content:
                {article_content[:5000]}
                """
                
                summary_response = client.chat.completions.create(
                    messages=[
                        {"role": "system", "content": "You are a helpful assistant that summarizes Wikipedia articles."},
                        {"role": "user", "content": summary_prompt}
                    ],
                    model="llama3-70b-8192",
                    temperature=0.3,
                    max_tokens=500
                )
                
                summary = summary_response.choices[0].message.content
            except Exception as e:
                summary = f"<p>Error generating summary: {str(e)}</p>"
            
            return {
                "title": article_title,
                "content": formatted_content,
                "summary": summary,
                "images": image_urls,
                "headings": headings
            }
        finally:
            driver.quit()
    except Exception as e:
        return {"error": str(e)}

# Document Proofreader utility
def proofread_document(file_path):
    try:
        client = get_groq_client()
        
        # Read the document content
        try:
            # Get file extension to determine document type
            file_extension = os.path.splitext(file_path)[1].lower()
            
            # Handle different file types appropriately
            if file_extension == '.pdf':
                try:
                    from PyPDF2 import PdfReader
                    reader = PdfReader(file_path)
                    document_content = '\n'.join([page.extract_text() for page in reader.pages if page.extract_text()])
                    if not document_content.strip():
                        return "<div class='error-message'>Could not extract text from PDF file. The file may be scanned or contain only images.</div>"
                except ImportError:
                    return "<div class='error-message'>PyPDF2 library not installed for PDF processing</div>"
                except Exception as e:
                    return f"<div class='error-message'>Error reading PDF: {str(e)}</div>"
            # Handle DOCX files
            elif file_extension in ['.docx', '.doc']:
                try:
                    import docx
                    doc = docx.Document(file_path)
                    document_content = '\n'.join([para.text for para in doc.paragraphs])
                except ImportError:
                    return "<div class='error-message'>python-docx library not installed for DOCX processing</div>"
                except Exception as e:
                    return f"<div class='error-message'>Error reading DOCX: {str(e)}</div>"
            # Handle text files with proper encoding detection
            else:
                try:
                    # First detect the file encoding
                    with open(file_path, 'rb') as file:
                        raw_data = file.read(10000)  # Read a sample to detect encoding
                        result = chardet.detect(raw_data)
                        encoding = result['encoding'] if result['encoding'] else 'utf-8'
                        confidence = result.get('confidence', 0)
                        
                        # If confidence is low, use a more reliable encoding
                        if confidence < 0.7:
                            encoding = 'utf-8'
                    
                    # Try to read with detected encoding
                    try:
                        with open(file_path, 'r', encoding=encoding) as file:
                            document_content = file.read()
                    except UnicodeDecodeError:
                        # First fallback: utf-8 with error replacement
                        try:
                            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                                document_content = file.read()
                        except UnicodeDecodeError:
                            # Second fallback: latin-1 (should handle any byte sequence)
                            with open(file_path, 'r', encoding='latin-1') as file:
                                document_content = file.read()
                except Exception as e:
                    return f"<div class='error-message'>Error reading file: {str(e)}</div>"
            
            # Get file extension to determine document type
            file_extension = os.path.splitext(file_path)[1].lower()
            document_type = "text"
            if file_extension in ['.md', '.markdown']:
                document_type = "markdown"
            elif file_extension in ['.tex']:
                document_type = "LaTeX"
            elif file_extension in ['.html', '.htm']:
                document_type = "HTML"
            
            # Create a system prompt for the proofreading
            system_prompt = """You are Myne, an advanced AI document proofreader assistant.
            Your task is to provide comprehensive analysis and suggestions for improving the document.
            Focus on grammar, spelling, tone, style, clarity, and coherence.
            Provide your feedback in a structured HTML format with clear sections and highlighting of issues."""
            
            # Create a prompt for the proofreading
            proofread_prompt = f"""
            Please perform a detailed analysis of the following {document_type} document:
            
            {document_content[:15000]}  # Increased content size limit
            
            Provide your analysis in the following format:
            
            1. SUMMARY: A brief overview of the document and its main issues
            2. GRAMMAR & SPELLING: List all grammar and spelling errors with corrections
            3. STYLE & TONE: Analyze the writing style and tone, suggesting improvements
            4. STRUCTURE & COHERENCE: Evaluate the document's structure and flow
            5. READABILITY: Assess the document's readability and suggest improvements
            6. ENHANCED VERSION: Provide an improved version of the document
            
            Format your response in HTML with appropriate headings, lists, and highlighting of issues.
            """
            
            response = client.chat.completions.create(
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": proofread_prompt}
                ],
                model="llama3-70b-8192",
                temperature=0.3,
                max_tokens=3000
            )
            
            proofreading_result = response.choices[0].message.content
            
            # Format the result with HTML for better display
            formatted_result = f"""
            <div class="proofreading-result">
                {proofreading_result}
            </div>
            """
            
            return formatted_result
        except Exception as e:
            return f"<div class='error-message'>Error proofreading document: {str(e)}</div>"
    except ValueError as e:
        # Handle missing API key
        return f"<div class='error-message'>Configuration Error: {str(e)}</div>"
    except Exception as e:
        # Handle other errors
        return f"<div class='error-message'>Sorry, I encountered an error: {str(e)}</div>"

# File handling utility
def handle_uploaded_file(file):
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    if file.size > MAX_FILE_SIZE:
        raise ValueError(f"File size exceeds maximum limit of {MAX_FILE_SIZE/1024/1024}MB")
    
    # Save file first
    file_path = default_storage.save(f'uploads/{file.name}', ContentFile(file.read()))
    full_path = os.path.join(settings.MEDIA_ROOT, file_path)
    
    try:
        # Determine file type
        mime_type, _ = mimetypes.guess_type(file.name)
        file_extension = os.path.splitext(file.name)[1].lower()
        
        # Handle text files (TXT, CSV, etc.)
        if mime_type == 'text/plain' or file_extension in ['.txt', '.csv', '.log', '.md']:
            # Detect encoding first to avoid codec errors
            with open(full_path, 'rb') as f:
                raw_data = f.read()
                result = chardet.detect(raw_data)
                encoding = result['encoding'] if result['encoding'] else 'utf-8'
                confidence = result.get('confidence', 0)
                
                # If confidence is low, use a more reliable encoding
                if confidence < 0.7:
                    encoding = 'utf-8'
            
            # Read with detected encoding, use multiple fallbacks if needed
            try:
                with open(full_path, 'r', encoding=encoding) as f:
                    content = f.read()
            except UnicodeDecodeError:
                # First fallback: utf-8 with error replacement
                try:
                    with open(full_path, 'r', encoding='utf-8', errors='replace') as f:
                        content = f.read()
                except UnicodeDecodeError:
                    # Second fallback: latin-1 (should handle any byte sequence)
                    with open(full_path, 'r', encoding='latin-1') as f:
                        content = f.read()
                        
            return {'path': full_path, 'content': content, 'type': 'text'}
            
        # Handle PDF files
        elif mime_type == 'application/pdf' or file_extension == '.pdf':
            try:
                # Import PyPDF2 for PDF processing
                from PyPDF2 import PdfReader
                # Return path only, content will be extracted when needed
                return {'path': full_path, 'type': 'pdf'}
            except ImportError:
                return {'path': full_path, 'type': 'document', 
                        'error': 'PyPDF2 library not installed for PDF processing'}
                
        # Handle DOCX files
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or \
             file_extension in ['.docx', '.doc']:
            try:
                # Import docx for DOCX processing
                import docx
                # Return path only, content will be extracted when needed
                return {'path': full_path, 'type': 'docx'}
            except ImportError:
                return {'path': full_path, 'type': 'document', 
                        'error': 'python-docx library not installed for DOCX processing'}
        else:
            raise ValueError(f"Unsupported file type: {mime_type or file_extension}")
    except Exception as e:
        # Clean up the file if there was an error
        if os.path.exists(full_path):
            os.remove(full_path)
        raise ValueError(f"Error processing file: {str(e)}")